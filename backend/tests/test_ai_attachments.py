"""Unit tests for tiqora.ai.attachments (text extraction + caps).

No DB/testcontainers needed — all fixtures are built in-memory with the same
libraries the extractor uses (docx/xlsx/odt roundtrip via python-docx/
openpyxl/odfpy; the PDF path is exercised via monkeypatching pypdf since that
library can only read, not author, PDFs).
"""

from __future__ import annotations

import io

import pytest

from tiqora.ai import attachments as attachments_mod
from tiqora.ai.attachments import (
    DEFAULT_RUN_BUDGET_CHARS,
    MAX_INPUT_BYTES,
    MAX_TEXT_CHARS,
    apply_budget,
    extract_attachment_text,
    is_image,
)

# ---------------------------------------------------------------------------
# is_image
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("content_type", "filename", "expected"),
    [
        ("image/png", "photo.png", True),
        ("image/jpeg", None, True),
        (None, "scan.JPG", True),
        ("application/pdf", "invoice.pdf", False),
        (None, "notes.txt", False),
        ("image/svg+xml", "diagram.svg", False),
    ],
)
def test_is_image(content_type: str | None, filename: str | None, expected: bool) -> None:
    assert is_image(content_type, filename) is expected


# ---------------------------------------------------------------------------
# Plaintext / CSV / HTML
# ---------------------------------------------------------------------------


def test_extract_plaintext() -> None:
    text = extract_attachment_text("notes.txt", "text/plain", b"Hello\nWorld")
    assert text == "Hello\nWorld"


def test_extract_csv() -> None:
    content = b"name,age\nAlice,30\nBob,40"
    text = extract_attachment_text("data.csv", "text/csv", content)
    assert text is not None
    assert "name | age" in text
    assert "Alice | 30" in text


def test_extract_html_strips_tags() -> None:
    content = b"<html><body><p>Hello <b>World</b></p></body></html>"
    text = extract_attachment_text("page.html", "text/html", content)
    assert text is not None
    assert "Hello" in text
    assert "<p>" not in text


def test_extract_unknown_type_returns_none() -> None:
    assert extract_attachment_text("archive.zip", "application/zip", b"PK\x03\x04") is None


# ---------------------------------------------------------------------------
# docx / xlsx / odt roundtrip (using the same libraries the extractor uses)
# ---------------------------------------------------------------------------


def test_extract_docx_roundtrip() -> None:
    import docx

    document = docx.Document()
    document.add_paragraph("Rechnung Nr. 4711")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Posten"
    table.rows[0].cells[1].text = "Betrag"
    buf = io.BytesIO()
    document.save(buf)

    text = extract_attachment_text(
        "rechnung.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buf.getvalue(),
    )
    assert text is not None
    assert "Rechnung Nr. 4711" in text
    assert "Posten | Betrag" in text


def test_extract_xlsx_roundtrip() -> None:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Summe"
    sheet.append(["Artikel", "Preis"])
    sheet.append(["Schraube", 1.5])
    buf = io.BytesIO()
    workbook.save(buf)

    text = extract_attachment_text(
        "tabelle.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        buf.getvalue(),
    )
    assert text is not None
    assert "[Sheet: Summe]" in text
    assert "Artikel | Preis" in text
    assert "Schraube | 1.5" in text


def test_extract_odt_roundtrip() -> None:
    from odf.opendocument import OpenDocumentText
    from odf.text import P

    doc = OpenDocumentText()
    doc.text.addElement(P(text="Sehr geehrte Damen und Herren"))
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_attachment_text(
        "brief.odt", "application/vnd.oasis.opendocument.text", buf.getvalue()
    )
    assert text is not None
    assert "Sehr geehrte Damen und Herren" in text


# ---------------------------------------------------------------------------
# PDF (dispatcher exercised via monkeypatch — pypdf cannot author PDFs)
# ---------------------------------------------------------------------------


def test_extract_pdf_dispatches_by_extension(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(attachments_mod, "_extract_pdf", lambda content: "PDF page text")
    text = extract_attachment_text("invoice.pdf", "application/pdf", b"%PDF-1.4 ...")
    assert text == "PDF page text"


def test_broken_pdf_returns_none_without_raising() -> None:
    # Not a real PDF — pypdf will raise internally; extractor must swallow it.
    assert extract_attachment_text("invoice.pdf", "application/pdf", b"not a pdf") is None


# ---------------------------------------------------------------------------
# Caps
# ---------------------------------------------------------------------------


def test_oversized_input_is_skipped() -> None:
    content = b"x" * (MAX_INPUT_BYTES + 1)
    assert extract_attachment_text("big.txt", "text/plain", content) is None


def test_text_is_truncated_with_marker() -> None:
    content = ("a" * (MAX_TEXT_CHARS + 500)).encode()
    text = extract_attachment_text("long.txt", "text/plain", content)
    assert text is not None
    assert len(text) <= MAX_TEXT_CHARS + len("\n… [gekürzt]")
    assert text.endswith("… [gekürzt]")


def test_empty_extracted_text_returns_none() -> None:
    assert extract_attachment_text("empty.txt", "text/plain", b"   \n  ") is None


def test_apply_budget_caps_total_and_marks_skipped() -> None:
    items = [("a", "x" * 30_000), ("b", "y" * 30_000), ("c", "z" * 10)]
    budgeted = apply_budget(items, budget_chars=50_000)
    assert budgeted[0] == ("a", "x" * 30_000)
    # second item is truncated to the remaining 20,000 chars + marker
    assert budgeted[1][0] == "b"
    assert len(budgeted[1][1]) <= 20_000 + len("\n… [gekürzt]")
    assert budgeted[1][1].endswith("… [gekürzt]")
    # third item arrives after the budget is exhausted
    assert budgeted[2] == ("c", "[Anhang übersprungen: budget]")


def test_apply_budget_default_is_50k() -> None:
    assert DEFAULT_RUN_BUDGET_CHARS == 50_000


def test_mask_attachment_block_keeps_labels_verbatim() -> None:
    """The summary prompts key on the literal '[Anhang: …]' label shape — a
    known name that happens to match inside the label (e.g. NER once tagged
    the capitalized noun 'Anhang' itself as a person) must not rewrite it."""
    from tiqora.ai.attachment_context import mask_attachment_block
    from tiqora.ai.pii import PiiMapper

    pii = PiiMapper(known_names=["Anhang", "Marie Knoblich"])
    block = (
        "[Anhang: 2.SK_Anhang_2_Aenderungsantraege.pdf — ca. 9396 Zeichen]\n"
        "Marie Knoblich stellt den Antrag, siehe Anhang.\n"
        "Kontakt: marie@example.com\n"
        "[Bild-Anhang: foto.jpg — Beschreibung durch Vision-Modell]\n"
        "Ein Foto von Marie Knoblich.\n"
        "[Anhang übersprungen: budget]"
    )
    masked = mask_attachment_block(pii, block)
    lines = masked.splitlines()
    assert lines[0] == "[Anhang: 2.SK_Anhang_2_Aenderungsantraege.pdf — ca. 9396 Zeichen]"
    assert lines[3] == "[Bild-Anhang: foto.jpg — Beschreibung durch Vision-Modell]"
    assert lines[5] == "[Anhang übersprungen: budget]"
    # Content between labels IS masked.
    assert "Marie Knoblich" not in lines[1]
    assert "[NAME_" in lines[1]
    assert "marie@example.com" not in masked
    assert "[EMAIL_" in masked
    assert "Marie Knoblich" not in lines[4]
