"""Text extraction for AI-context document attachments (not images — see
:mod:`tiqora.ai.vision` for the image-description pre-pass).

Extracted text is embedded directly into the LLM context of the ticket
render (``[Anhang: rechnung.pdf]\\n<text>``), unlike images, which are never
shown to the main model. Every parser call is wrapped so a corrupt/unusual
attachment can never abort an agent run — worst case, the attachment is
silently skipped (with a structlog warning).

Caps (deliberately generous but bounded):

- Input larger than :data:`MAX_INPUT_BYTES` (10 MB) is not parsed at all.
- Extracted text per attachment is capped at :data:`MAX_TEXT_CHARS`
  (20,000 chars), with a trailing "… [gekürzt]" marker if truncated.
- The caller enforces a total per-run budget across all attachments
  (:data:`DEFAULT_RUN_BUDGET_CHARS`) — see :func:`apply_budget`.
"""

from __future__ import annotations

import csv
import io

import structlog

logger = structlog.get_logger(__name__)

MAX_INPUT_BYTES = 10 * 1024 * 1024
MAX_TEXT_CHARS = 20_000
DEFAULT_RUN_BUDGET_CHARS = 50_000

_TRUNCATION_MARKER = "… [gekürzt]"
_BUDGET_SKIPPED_MARKER = "[Anhang übersprungen: budget]"

_TEXT_CONTENT_TYPES = {"text/plain", "text/csv", "text/markdown"}
_HTML_CONTENT_TYPES = {"text/html"}
_IMAGE_CONTENT_TYPES = {"image/png", "image/jpeg", "image/gif", "image/webp"}

_TEXT_EXTENSIONS = (".txt", ".csv", ".md", ".markdown", ".log")
_HTML_EXTENSIONS = (".html", ".htm")


def _cap_text(text: str) -> str:
    if len(text) <= MAX_TEXT_CHARS:
        return text
    return text[:MAX_TEXT_CHARS].rstrip() + "\n" + _TRUNCATION_MARKER


def _extract_pdf(content: bytes) -> str | None:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p for p in pages if p.strip())


def _extract_docx(content: bytes) -> str | None:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(content: bytes) -> str | None:
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    try:
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"[Sheet: {sheet.title}]")
            for row_idx, row in enumerate(sheet.iter_rows(values_only=True)):
                if row_idx >= 200:
                    parts.append("… [Sheet gekürzt]")
                    break
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    finally:
        workbook.close()


def _extract_odf(content: bytes) -> str | None:
    from odf import opendocument, table
    from odf import text as odf_text

    doc = opendocument.load(io.BytesIO(content))
    parts: list[str] = []
    for paragraph in doc.getElementsByType(odf_text.P):
        value = str(paragraph)
        if value.strip():
            parts.append(value)
    for tbl in doc.getElementsByType(table.Table):
        for row in tbl.getElementsByType(table.TableRow):
            cells = []
            for cell in row.getElementsByType(table.TableCell):
                cells.append(str(cell))
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_plaintext(content: bytes) -> str | None:
    return content.decode("utf-8", errors="replace")


def _extract_html(content: bytes) -> str | None:
    from tiqora.domain.quoting import html_to_plaintext

    return html_to_plaintext(content.decode("utf-8", errors="replace"))


def _extract_csv(content: bytes) -> str | None:
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    lines = [" | ".join(row) for row in reader]
    return "\n".join(lines)


def is_image(content_type: str | None, filename: str | None) -> bool:
    """True if this attachment should go through the vision pre-pass instead
    of text extraction (see module docstring)."""
    ct = (content_type or "").split(";")[0].strip().lower()
    if ct in _IMAGE_CONTENT_TYPES:
        return True
    name = (filename or "").lower()
    return name.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp"))


def extract_attachment_text(
    filename: str | None, content_type: str | None, content: bytes
) -> str | None:
    """Dispatch to a format-specific extractor by content type / filename
    extension. Returns ``None`` for unknown/unsupported types, oversized
    input, or any parser failure — never raises.
    """
    if len(content) > MAX_INPUT_BYTES:
        logger.warning(
            "ai_attachment_extract_skipped_too_large", filename=filename, size=len(content)
        )
        return None

    ct = (content_type or "").split(";")[0].strip().lower()
    name = (filename or "").lower()

    try:
        text: str | None
        if ct == "application/pdf" or name.endswith(".pdf"):
            text = _extract_pdf(content)
        elif (
            ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or name.endswith(".docx")
        ):
            text = _extract_docx(content)
        elif (
            ct == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            or name.endswith(".xlsx")
        ):
            text = _extract_xlsx(content)
        elif ct in (
            "application/vnd.oasis.opendocument.text",
            "application/vnd.oasis.opendocument.spreadsheet",
        ) or name.endswith((".odt", ".ods")):
            text = _extract_odf(content)
        elif ct == "text/csv" or name.endswith(".csv"):
            text = _extract_csv(content)
        elif ct in _HTML_CONTENT_TYPES or name.endswith(_HTML_EXTENSIONS):
            text = _extract_html(content)
        elif ct in _TEXT_CONTENT_TYPES or ct.startswith("text/") or name.endswith(_TEXT_EXTENSIONS):
            text = _extract_plaintext(content)
        else:
            return None
    except Exception:  # noqa: BLE001 — a broken/unusual file must never break a run
        logger.warning(
            "ai_attachment_extract_failed",
            filename=filename,
            content_type=content_type,
            exc_info=True,
        )
        return None

    if text is None or not text.strip():
        return None
    return _cap_text(text)


def apply_budget(
    items: list[tuple[str, str]], *, budget_chars: int = DEFAULT_RUN_BUDGET_CHARS
) -> list[tuple[str, str]]:
    """Apply a total per-run character budget across already-extracted
    ``(label, text)`` attachment texts, in the given (chronological) order.
    Once the budget is exhausted, remaining attachments are replaced with a
    "budget" skip marker rather than dropped silently.
    """
    remaining = budget_chars
    out: list[tuple[str, str]] = []
    for label, text in items:
        if remaining <= 0:
            out.append((label, _BUDGET_SKIPPED_MARKER))
            continue
        if len(text) > remaining:
            text = text[:remaining].rstrip() + "\n" + _TRUNCATION_MARKER
        remaining -= len(text)
        out.append((label, text))
    return out


__all__ = [
    "DEFAULT_RUN_BUDGET_CHARS",
    "MAX_INPUT_BYTES",
    "MAX_TEXT_CHARS",
    "apply_budget",
    "extract_attachment_text",
    "is_image",
]
